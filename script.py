#import modules
import pandas as pd
import numpy as np
import seaborn as sns
import matplotlib.pyplot as plt
from docx import Document
import os

#validate files are present
if not os.path.exists("set_3/A_vs_B.deseq2.results.tsv") or not os.path.exists("set_3/A_vs_E.deseq2.results.tsv"):
    raise FileNotFoundError("One or more files are missing or misnamed")

#open file to record outputs to
output_file = Document()
output_file.save("Output_file.docx")

if not os.path.exists("Outputs"):
    os.mkdir("Outputs")

#Read in data
AvsB = pd.read_csv("set_3/A_vs_B.deseq2.results.tsv", sep="\t")
AvsE = pd.read_csv("set_3/A_vs_E.deseq2.results.tsv", sep="\t")

#validate required columns exist
required_columns = {"gene_id", "baseMean", "log2FoldChange", "pvalue", "padj"}
if not required_columns.issubset(AvsB.columns):
    raise ValueError(f"AvsB is missing required columns: {required_columns - set(AvsB.columns)}")

if not required_columns.issubset(AvsE.columns):
    raise ValueError(f"AvsE is missing required columns: {required_columns - set(AvsE.columns)}")

#Filter insignificant results
filtered_AvsB = AvsB[(AvsB["log2FoldChange"].abs() >= 2) & (AvsB["pvalue"]<0.05)]
filtered_AvsE = AvsE[(AvsE["log2FoldChange"].abs() >= 2) & (AvsE["pvalue"]<0.05)]

#Export filtered results as csv
filtered_AvsB[["gene_id", "log2FoldChange", "pvalue", "padj"]].to_csv("Outputs/filtered_AvsB.csv", index = False)
filtered_AvsE[["gene_id", "log2FoldChange", "pvalue", "padj"]].to_csv("Outputs/filtered_AvsE.csv", index = False)

#Filter for upregulated genes
AvsB_upreg = filtered_AvsB[filtered_AvsB["log2FoldChange"]>0]
AvsE_upreg = filtered_AvsE[filtered_AvsE["log2FoldChange"]>0]

#Filter for downregulated genes
AvsB_downreg = filtered_AvsB[filtered_AvsB["log2FoldChange"]<0]
AvsE_downreg = filtered_AvsE[filtered_AvsE["log2FoldChange"]<0]

#output results
output_file.add_heading("AvB Statistical summary") #add heading to stats

output_file.add_paragraph("\n Upregulated genes: " + str(AvsB_upreg.shape[0]) + #AvB upregulated genes
                  "\n Downregulated genes: " + str(AvsB_downreg.shape[0]) + #AvB downregulated genes
                  "\n \n Maximum p value: " + str(max(AvsB["pvalue"])) + #AvB maximum p value
                  "\n Minimum p value: " + str(min(AvsB["pvalue"])) + #AvB minimum p value
                  "\n Median p value: " + str(round(AvsB["pvalue"].median(), 2)) + #AvB median p value
                  "\n Standard deviation of p value: " + str(round(AvsB["pvalue"].std(), 2)) + #AvB p value standard deviation
                  "\n \n Maximum 2 fold log: " + str(max(AvsB["log2FoldChange"])) + #AvB maximum 2 fold log
                  "\n Minimum 2 fold log: " + str(min(AvsB["log2FoldChange"])) + #AvB minimum 2 fold log
                  "\n Median 2 fold log: " + str(round(AvsB["log2FoldChange"].median(), 2)) + #AvB median 2 fold log
                  "\n Standard deviation of 2 fold log: " + str(round(AvsB["log2FoldChange"].std(), 2)) + "\n \n") #AvB standard deviation 2 fold log

output_file.add_heading("AvE Statistical summary") #add heading to stats

output_file.add_paragraph("\n Upregulated genes: " + str(AvsE_upreg.shape[0]) + #AvE upregulated genes
                  "\n Downregulated genes: " + str(AvsE_downreg.shape[0]) + #AvE downregulated genes
                  "\n \n Maximum p value: " + str(max(AvsE["pvalue"])) + #AvE maximum p value
                  "\n Minimum p value: " + str(min(AvsE["pvalue"])) + #AvE minimum p value
                  "\n Median p value: " + str(round(AvsE["pvalue"].median(), 2)) + #AvE median p value
                  "\n Standard deviation of p value: " + str(round(AvsE["pvalue"].std(), 2)) + #AvE p value standard deviation
                  "\n \n Maximum 2 fold log: " + str(max(AvsE["log2FoldChange"])) + #AvE maximum 2 fold log
                  "\n Minimum 2 fold log: " + str(min(AvsE["log2FoldChange"])) + #AvE minimum 2 fold log
                  "\n Median 2 fold log: " + str(round(AvsE["log2FoldChange"].median(), 2)) + #AvE median 2 fold log
                  "\n Standard deviation of 2 fold log: " + str(round(AvsE["log2FoldChange"].std(), 2))) #AvE standard deviation 2 fold log

#generate negative log values for pvalues
AvsB["pvalue_neg_log10"] = -np.log10(AvsB["pvalue"])
AvsE["pvalue_neg_log10"] = -np.log10(AvsE["pvalue"])

#Generate AvsB volcano graph/scatter plot
AvsB_volcano = sns.scatterplot(x = AvsB["log2FoldChange"], y = AvsB["pvalue_neg_log10"]) #set values for AvsB plot
AvsB_volcano.set(xlabel = "Magnitude of changes (log2)", ylabel = "Significance (negative log 10)", title = "AvsB Volcano graph") #set labels for AvsB plot
plt.savefig("Outputs/AvsB_volcano.png") #save plot
plt.close() #close plot within python
output_file.add_heading("AvsB Volcano plot") #add heading to plot
output_file.add_picture("Outputs/AvsB_volcano.png") #add picture to output document

#Generate AvsE volcano graph/scatter plot
AvsE_volcano = sns.scatterplot(x = AvsE["log2FoldChange"], y = AvsE["pvalue_neg_log10"]) #set values for AvsE plot
AvsE_volcano.set(xlabel = "Magnitude of changes (log2)", ylabel = "Significance (negative log 10)", title = "AvsE Volcano graph") #set labels for AvsE plot
plt.savefig("Outputs/AvsE_volcano.png") #save plot
plt.close() #close plot
output_file.add_heading("AvsE Volcano plot") #add heading to plot
output_file.add_picture("Outputs/AvsE_volcano.png") #add picture to output document

#AvsB MA plots
plt.scatter(x = np.log2(AvsB["baseMean"]), y = AvsB["log2FoldChange"], s=5) #set baseMean to log2, and use as x axis.
plt.xlabel("Log2 of base mean") #set labels
plt.ylabel("Log2 of fold change")
plt.title("MA plot for AvsB data")
plt.savefig("Outputs/AvsB_MA_plot.png") #save the plot
plt.close() #close the plot
output_file.add_heading("AvsB MA plot") #add heading to plot
output_file.add_picture("Outputs/AvsB_MA_plot.png") #add picture to output document


#AvsE MA plots
plt.scatter(x = np.log2(AvsE["baseMean"]), y = AvsE["log2FoldChange"], s=5) #set baseMean to log2, and use as x axis.
plt.xlabel("Log2 of base mean") #set labels
plt.ylabel("Log2 of fold change")
plt.title("MA plot for AvsE data")
plt.savefig("Outputs/AvsE_MA_plot.png")
plt.close() #close the plot
output_file.add_heading("AvsE MA plot") #add heading to plot
output_file.add_picture("Outputs/AvsE_MA_plot.png") #add picture to output document

#AvsB p value histogram
plt.hist(AvsB["pvalue"], bins = 15) #15 bins
plt.xlabel("P value") #add labels
plt.ylabel("Counts")
plt.title("Histogram of p values across AvsB data")
plt.savefig("Outputs/AvsB_pvalue_histogram.png")
plt.close()
output_file.add_heading("AvsB p value histogram") #add heading to plot
output_file.add_picture("Outputs/AvsB_pvalue_histogram.png") #add picture to output document

#AvsE pvalue histogram
plt.hist(AvsE["pvalue"], bins = 15) #15 bins
plt.xlabel("P value") #add labels
plt.ylabel("Counts")
plt.title("Histogram of p values across AvsE data")
plt.savefig("Outputs/AvsE_pvalue_histogram.png")
plt.close()
output_file.add_heading("AvsB p value histogram") #add heading to plot
output_file.add_picture("Outputs/AvsE_pvalue_histogram.png") #add picture to output document

#AvsB Heatmap
arranged_AvsB = AvsB.sort_values("log2FoldChange", ascending = False) #sort data in descending order
arranged_AvsB["log2_baseMean"] = np.log2(AvsB["baseMean"]) #add log2 column
top_10_AvsB = arranged_AvsB.head(10) #pick the top 10 highest log change results
AvsB_heatmap = top_10_AvsB[["log2FoldChange", "log2_baseMean"]] #select relevant columns

sns.heatmap(AvsB_heatmap, annot = True, cmap = "rocket")
plt.ylabel("Genes of interest")
plt.title("Heat map showing the log values of the \nBase Mean and Fold Change for the top 10 genes of AvsB")
plt.savefig("Outputs/Heat_map_for_MeanvsFoldChange_AvsB.png")
plt.close()
output_file.add_heading("AvsB heat map") #add heading to plot
output_file.add_picture("Outputs/Heat_map_for_MeanvsFoldChange_AvsB.png") #add picture to output document

#AvsE Heatmap
arranged_AvsE = AvsE.sort_values("log2FoldChange", ascending = False) #sort data in descending order
arranged_AvsE["log2_baseMean"] = np.log2(AvsE["baseMean"]) #add log2 column
top_10_AvsE = arranged_AvsE.head(10) #pick the top 10 highest log change results
AvsE_heatmap = top_10_AvsE[["log2FoldChange", "log2_baseMean"]] #select relevant columns

sns.heatmap(AvsE_heatmap, annot = True, cmap = "rocket")
plt.ylabel("Genes of interest")
plt.title("Heat map showing the log values of the \nBase Mean and Fold Change for the top 10 genes of AvsE")
plt.savefig("Outputs/Heat_map_for_MeanvsFoldChange_AvsE.png")
plt.close()
output_file.add_heading("AvsE heat map") #add heading to plot
output_file.add_picture("Outputs/Heat_map_for_MeanvsFoldChange_AvsE.png") #add picture to output document

define_significance_header = output_file.add_heading("Defining Significance")

define_significance = output_file.add_paragraph("\n Significance was defined as p<0.05, and a log fold change of greater than 2 or less than -2 \n")

# Convert filtered AvsB data to list of lists
filtered_AvsB_table_data = [filtered_AvsB.columns.to_list()] + filtered_AvsB.values.tolist()

# Add a heading for the table
output_file.add_heading("Filtered AvsB Table")

# Add table to output document
AvsB_filtered_table = output_file.add_table(rows=0, cols=len(filtered_AvsB.columns))

#add data to table
for row_data in filtered_AvsB_table_data:
    row_cells = AvsB_filtered_table.add_row().cells
    for idx, value in enumerate(row_data):
        row_cells[idx].text = str(value)

# Convert filtered AvsE data to list of lists
filtered_AvsE_table_data = [filtered_AvsE.columns.to_list()] + filtered_AvsE.values.tolist()

# Add a heading for the table
output_file.add_heading("Filtered AvsE Table")

# Add table to output document
AvsB_filtered_table = output_file.add_table(rows=0, cols=len(filtered_AvsE.columns))

#add data to table
for row_data in filtered_AvsE_table_data:
    row_cells = AvsB_filtered_table.add_row().cells
    for idx, value in enumerate(row_data):
        row_cells[idx].text = str(value)

#save word document
output_file.save("Output_file.docx")